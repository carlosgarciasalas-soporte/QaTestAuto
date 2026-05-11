from __future__ import annotations

import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "Informe_Consolidado_QA_EFAC.docx"


NAVY = RGBColor(31, 78, 121)
TEAL = RGBColor(15, 118, 110)
GRAY = RGBColor(100, 116, 139)
LIGHT_BLUE = "D9EAF7"
LIGHT_TEAL = "DDF3F0"
LIGHT_GRAY = "F2F4F7"


def git_value(*args: str) -> str:
    try:
        return subprocess.check_output(["git", *args], cwd=ROOT, text=True).strip()
    except Exception:
        return "No disponible"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D0D7DE")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, color=NAVY)
        set_cell_shading(cell, LIGHT_BLUE)

    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            set_cell_text(cell, value)
            if index == 0:
                set_cell_shading(cell, LIGHT_GRAY)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)

    document.add_paragraph()
    return table


def add_callout(document: Document, title: str, body: str, fill: str = LIGHT_TEAL) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = TEAL
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 24, NAVY),
        ("Heading 1", 17, NAVY),
        ("Heading 2", 13, TEAL),
        ("Heading 3", 11, GRAY),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def set_document_layout(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)

    header = section.header.paragraphs[0]
    header.text = "EFAC - Informe consolidado QA"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = GRAY

    footer = section.footer.paragraphs[0]
    footer.text = "Proyecto academico de pruebas y calidad del software"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8.5)
    footer.runs[0].font.color.rgb = GRAY


def add_cover(document: Document, commit: str, branch: str) -> None:
    for _ in range(2):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Informe Consolidado QA\nSistema CRUD de Clientes EFAC")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Analisis del proyecto, estrategia de pruebas, automatizacion y evidencias")
    run.font.size = Pt(13)
    run.font.color.rgb = TEAL

    document.add_paragraph()
    metadata = [
        ["Proyecto", "EFAC QA Testing y Automatizacion"],
        ["Tecnologia", "ASP.NET Core .NET 9, Razor Pages, Web API, xUnit, Selenium"],
        ["Repositorio", "https://github.com/carlosgarciasalas-soporte/QaTestAuto.git"],
        ["Rama / Commit", f"{branch} / {commit}"],
        ["Fecha de consolidacion", datetime.now().strftime("%Y-%m-%d %H:%M")],
    ]
    add_table(document, ["Dato", "Valor"], metadata, [4.0, 12.0])

    add_callout(
        document,
        "Proposito del documento",
        "Consolidar en un solo informe la informacion relevante del README, la matriz de evidencias, "
        "el plan de automatizacion y el estado actual de pruebas automaticas del proyecto EFAC.",
    )
    document.add_page_break()


def build_document() -> Document:
    commit = git_value("rev-parse", "--short", "HEAD")
    branch = git_value("branch", "--show-current")

    document = Document()
    configure_styles(document)
    set_document_layout(document)
    add_cover(document, commit, branch)

    document.add_heading("1. Resumen ejecutivo", level=1)
    document.add_paragraph(
        "EFAC es un sistema CRUD de clientes construido con ASP.NET Core .NET 9. "
        "Su objetivo funcional es administrar clientes bajo reglas de negocio asociadas a cumplimiento DIAN, "
        "incluyendo normalizacion de NIT, calculo de digito de verificacion, control de duplicidad, "
        "validacion de mayoria de edad y diferenciacion entre persona natural y juridica."
    )
    add_callout(
        document,
        "Estado general",
        "El proyecto cuenta con pruebas automatizadas de API y UI, script de ejecucion completo para Windows, "
        "reportes TRX, log persistente y documentacion de evidencias. La ultima validacion registrada reporta "
        "11 pruebas API superadas y 10 pruebas Selenium superadas.",
    )

    add_table(
        document,
        ["Area", "Resultado consolidado"],
        [
            ["Backend/API", "Endpoints CRUD y calculo DV disponibles en /api/clientes."],
            ["Frontend", "Interfaz Razor Pages con IDs estables para automatizacion Selenium."],
            ["Persistencia", "Entity Framework Core InMemory para ejecucion academica y pruebas."],
            ["Automatizacion API", "11 pruebas xUnit con WebApplicationFactory y FluentAssertions."],
            ["Automatizacion UI", "10 pruebas Selenium con Chrome visible o headless."],
            ["Evidencias", "Log QA, reportes TRX, guia de capturas y matriz de trazabilidad."],
        ],
        [4.0, 12.0],
    )

    document.add_heading("2. Alcance y fuentes consolidadas", level=1)
    document.add_paragraph(
        "Este informe consolida la informacion operativa y tecnica necesaria para revisar el proyecto, "
        "analizar su cobertura de pruebas y ejecutar el ciclo QA automatizado."
    )
    add_table(
        document,
        ["Fuente", "Contenido integrado"],
        [
            ["README.md", "Descripcion, arquitectura, reglas, ejecucion, automatizacion y evidencias."],
            ["QA_Matriz_Evidencias_Pruebas.md", "Trazabilidad, casos manuales, evidencias, campos y resultados esperados."],
            ["AUTOMATIZACION_PLAN.md", "Etapas de automatizacion, criterios de avance y cobertura automatizada."],
            ["Codigo de pruebas", "Casos API y Selenium implementados en los proyectos de prueba."],
            ["Scripts QA", "Ejecucion automatica con run-qa-tests.bat y run-qa-tests.ps1."],
        ],
        [5.0, 11.0],
    )

    document.add_heading("3. Vision tecnica del proyecto", level=1)
    document.add_heading("3.1 Arquitectura", level=2)
    add_table(
        document,
        ["Proyecto", "Responsabilidad"],
        [
            ["Efac.Domain", "Entidad Cliente, enums, excepciones y calculo DIAN modulo 11."],
            ["Efac.Application", "DTOs, servicios de caso de uso, validaciones y contratos."],
            ["Efac.Infrastructure", "DbContext, repositorio e inyeccion de dependencias de persistencia."],
            ["Efac.WebAPI", "Controladores REST, Swagger, Razor Pages e interfaz web."],
            ["Efac.Tests.Api", "Pruebas automatizadas de endpoints y reglas de negocio."],
            ["Efac.Tests.Selenium", "Pruebas E2E de interfaz con navegador."],
            ["test-assets", "Estructura para evidencias de ejecucion, capturas y reportes."],
        ],
        [4.0, 12.0],
    )

    document.add_heading("3.2 Endpoints disponibles", level=2)
    add_table(
        document,
        ["Metodo", "Ruta", "Uso"],
        [
            ["GET", "/api/clientes", "Listar clientes."],
            ["GET", "/api/clientes/{id}", "Consultar cliente por identificador."],
            ["GET", "/api/clientes/calcular-dv/{nit}", "Normalizar NIT y calcular DV."],
            ["POST", "/api/clientes", "Crear cliente natural o juridico."],
            ["PUT", "/api/clientes/{id}", "Actualizar cliente existente."],
            ["DELETE", "/api/clientes/{id}", "Eliminar cliente existente."],
        ],
        [2.5, 6.0, 7.5],
    )

    document.add_heading("3.3 Reglas de negocio principales", level=2)
    add_bullets(
        document,
        [
            "El NIT se normaliza antes de persistir, buscar o calcular el DV.",
            "El digito de verificacion se calcula con algoritmo DIAN modulo 11.",
            "No se permite registrar dos clientes con el mismo NIT.",
            "La persona natural requiere nombres, apellidos y fecha de nacimiento.",
            "La persona natural menor de edad debe ser rechazada.",
            "La persona juridica requiere razon social.",
            "Los campos no aplicables al tipo de persona no deben persistirse.",
        ],
    )

    document.add_heading("4. Estrategia QA", level=1)
    add_table(
        document,
        ["Nivel", "Herramienta", "Objetivo"],
        [
            ["Prueba manual API", "Swagger", "Validar contratos y payloads durante exploracion manual."],
            ["Prueba automatizada API", "xUnit + WebApplicationFactory", "Validar reglas HTTP, negocio y persistencia en memoria."],
            ["Prueba E2E UI", "Selenium WebDriver", "Simular flujos reales de usuario en navegador."],
            ["Evidencias", "Logs, TRX y capturas", "Sustentar ejecucion y resultado de pruebas."],
            ["Orquestacion", "run-qa-tests.bat / .ps1", "Ejecutar restore, build, API, UI, suite completa y reportes."],
        ],
        [4.0, 5.0, 7.0],
    )

    document.add_heading("5. Cobertura automatizada", level=1)
    document.add_heading("5.1 Pruebas API", level=2)
    add_table(
        document,
        ["ID", "Caso", "Resultado esperado"],
        [
            ["API-001", "Listar clientes", "200 OK con datos semilla."],
            ["API-002", "Consultar cliente inexistente", "404 Not Found."],
            ["API-003", "Calcular DV con NIT formateado", "200 OK, NIT normalizado y DV calculado."],
            ["API-004", "Rechazar NIT invalido", "400 Bad Request."],
            ["API-005", "Crear persona natural valida", "201 Created."],
            ["API-006", "Rechazar NIT duplicado", "409 Conflict."],
            ["API-007", "Rechazar menor de edad", "400 Bad Request."],
            ["API-008", "Rechazar natural sin fecha", "400 Bad Request."],
            ["API-009", "Rechazar juridica sin razon social", "400 Bad Request."],
            ["API-010", "Actualizar cliente", "200 OK con datos actualizados."],
            ["API-011", "Eliminar cliente", "204 No Content y posterior 404."],
        ],
        [2.5, 7.0, 6.5],
    )

    document.add_heading("5.2 Pruebas Selenium", level=2)
    add_table(
        document,
        ["ID", "Caso", "Controles o flujo validado"],
        [
            ["UI-001", "Carga principal", "Titulo, input-search y btn-new-client."],
            ["UI-002", "Busqueda por NIT", "input-search y tabla de clientes."],
            ["UI-003", "Alternancia Natural/Juridica", "Campos naturales y razon social."],
            ["UI-004", "Calculo DV", "input-nit, input-dv y llamada a API."],
            ["UI-005", "Crear natural", "Formulario completo y aparicion en tabla."],
            ["UI-006", "Crear juridica", "Tipo persona, razon social y guardado."],
            ["UI-007", "NIT duplicado", "form-validation-summary con error esperado."],
            ["UI-008", "Menor de edad", "fechaNacimiento y mensaje de error."],
            ["UI-009", "Editar cliente", "Accion Editar y persistencia del cambio."],
            ["UI-010", "Eliminar cliente", "Accion Eliminar, confirmacion y ausencia en tabla."],
        ],
        [2.5, 6.0, 7.5],
    )

    document.add_heading("5.3 Campos validados", level=2)
    add_table(
        document,
        ["Campo", "Validacion"],
        [
            ["tipoPersona", "Diferencia Natural y Juridica; activa/oculta campos segun tipo."],
            ["nit", "Normalizacion, busqueda, duplicidad y calculo de DV."],
            ["dv", "Calculo DIAN modulo 11."],
            ["nombres / apellidos", "Obligatorios y visibles para persona natural."],
            ["fechaNacimiento", "Obligatoria para natural; valida mayoria de edad."],
            ["razonSocial", "Obligatoria y visible para persona juridica."],
            ["email / telefono / direccion", "Persistencia en creacion y actualizacion."],
            ["ciudadCodigoMunicipio", "Persistencia del municipio."],
            ["responsabilidadFiscal", "Persistencia del regimen fiscal seleccionado."],
        ],
        [5.0, 11.0],
    )

    document.add_heading("6. Flujo de ejecucion y evidencias", level=1)
    document.add_heading("6.1 Ejecucion automatica", level=2)
    add_numbered(
        document,
        [
            "Abrir CMD o PowerShell en la raiz del proyecto.",
            "Ejecutar .\\run-qa-tests.bat.",
            "Verificar el encabezado EFAC - EJECUCION QA AUTOMATICA.",
            "Esperar restore, build, pruebas API, WebAPI local, Selenium y suite completa.",
            "Tomar capturas del navegador y consola durante la ejecucion.",
            "Revisar la carpeta indicada en Reportes TRX.",
            "Presionar una tecla para cerrar la ventana al finalizar.",
        ],
    )

    document.add_heading("6.2 Archivos generados", level=2)
    add_table(
        document,
        ["Archivo", "Evidencia"],
        [
            ["qa-execution.log", "Registro completo de comandos, etapas y resultados."],
            ["api-tests.trx", "Resultado formal de las 11 pruebas API."],
            ["selenium-tests.trx", "Resultado formal de las 10 pruebas Selenium."],
            ["*.trx adicionales", "Resultados de la suite completa."],
            ["webapi.log", "Salida de la WebAPI durante el ciclo automatico."],
        ],
        [5.0, 11.0],
    )

    document.add_heading("6.3 Capturas recomendadas", level=2)
    add_bullets(
        document,
        [
            "Inicio del script con el encabezado de ejecucion automatica.",
            "Compilacion correcta.",
            "Pruebas API con Superado: 11.",
            "Chrome visible llenando formularios, buscando, editando y eliminando clientes.",
            "Errores visibles por NIT duplicado y menor de edad.",
            "Pruebas Selenium con Superado: 10.",
            "Resumen final con Restore, Build, API, Selenium y Suite completa en OK.",
            "Carpeta generada en test-assets/evidence/reports.",
        ],
    )

    document.add_heading("7. Evidencia visible en GitHub", level=1)
    add_table(
        document,
        ["Tipo", "Archivo", "Evidencia"],
        [
            ["Codigo", "Efac.Tests.Api/ClientesApiTests.cs", "Casos API y reglas de negocio."],
            ["Codigo", "Efac.Tests.Selenium/Tests/ClientesUiSmokeTests.cs", "Flujos UI CRUD y validaciones negativas."],
            ["Codigo", "Efac.Tests.Selenium/Pages/ClientesPage.cs", "Page Object y acciones de navegador."],
            ["Codigo", "run-qa-tests.bat", "Entrada de ejecucion para Windows con pausa final."],
            ["Codigo", "run-qa-tests.ps1", "Orquestacion completa y generacion de evidencias."],
            ["Documentacion", "README.md", "Guia de proyecto, pruebas, ejecucion y evidencias."],
            ["Documentacion", "QA_Matriz_Evidencias_Pruebas.md", "Trazabilidad y checklist de evidencias."],
            ["Documentacion", "AUTOMATIZACION_PLAN.md", "Plan por etapas y cobertura."],
        ],
        [3.0, 6.0, 7.0],
    )

    document.add_heading("8. Analisis del proyecto", level=1)
    document.add_heading("8.1 Fortalezas", level=2)
    add_bullets(
        document,
        [
            "Separacion clara por capas y por proyectos de prueba.",
            "Cobertura API alineada con reglas de negocio criticas.",
            "Selenium visible permite evidencias demostrativas durante la entrega.",
            "El Page Object reduce fragilidad y centraliza selectores.",
            "El script automatico deja trazabilidad por ejecucion mediante logs y TRX.",
            "La documentacion conecta requisitos, casos, campos y evidencia.",
        ],
    )

    document.add_heading("8.2 Riesgos y consideraciones", level=2)
    add_bullets(
        document,
        [
            "La persistencia InMemory es adecuada para entorno academico, pero no reemplaza pruebas contra base de datos real.",
            "Las capturas deben tomarse manualmente durante ejecucion visible si se requieren como evidencia grafica final.",
            "Selenium visible depende de Chrome/ChromeDriver disponible en el equipo.",
            "Para despliegue productivo se recomienda una plataforma compatible con ASP.NET Core, no Vercel directamente.",
        ],
    )

    document.add_heading("8.3 Recomendaciones", level=2)
    add_bullets(
        document,
        [
            "Conservar el PDF, el README y la matriz como paquete de entrega.",
            "Ejecutar run-qa-tests.bat antes de cada entrega formal.",
            "Guardar capturas finales en test-assets/evidence/screenshots si el docente las solicita.",
            "Agregar reportes HTML en una etapa futura si se requiere lectura mas amigable que TRX.",
            "Separar frontend y backend si se desea publicar una version web en Vercel.",
        ],
    )

    document.add_heading("9. Conclusiones", level=1)
    document.add_paragraph(
        "El proyecto EFAC presenta una base tecnica consistente para pruebas y calidad de software. "
        "La automatizacion cubre escenarios funcionales clave de API y UI, incluyendo casos positivos, "
        "validaciones negativas, CRUD visual con Selenium y generacion de evidencias. "
        "La estructura documental y los scripts permiten reproducir el ciclo QA de forma clara, auditable y apta para entrega academica."
    )

    document.add_page_break()
    document.add_heading("Anexo A. Comandos principales", level=1)
    add_table(
        document,
        ["Objetivo", "Comando"],
        [
            ["Restaurar", "dotnet restore Efac.sln"],
            ["Compilar", "dotnet build Efac.sln --no-restore"],
            ["Pruebas API", "dotnet test Efac.Tests.Api --no-build"],
            ["Levantar WebAPI", "dotnet run --project Efac.WebAPI --launch-profile http"],
            ["Activar Selenium", "$env:EFAC_RUN_SELENIUM=\"true\""],
            ["Base URL Selenium", "$env:EFAC_BASE_URL=\"http://localhost:5100/\""],
            ["Chrome visible", "$env:EFAC_SELENIUM_HEADLESS=\"false\""],
            ["Pruebas Selenium", "dotnet test Efac.Tests.Selenium --no-build"],
            ["Ciclo completo", ".\\run-qa-tests.bat"],
        ],
        [4.0, 12.0],
    )

    document.add_heading("Anexo B. Resultado esperado de ejecucion", level=1)
    add_callout(
        document,
        "Resumen esperado",
        "Restore: OK | Build: OK | Pruebas API: OK | Pruebas Selenium: OK | Suite completa: OK",
        fill=LIGHT_BLUE,
    )

    return document


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    document = build_document()
    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
